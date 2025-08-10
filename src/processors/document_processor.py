"""
Document processor for handling PDF, Word, and image files.
"""
import io
import re
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
import pandas as pd

# Import logger first before trying to use it
from src.utils.logger import logger

# PDF processing
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    logger.warning("PyPDF2 not available. PDF processing will be limited.")
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    logger.warning("pdfplumber not available. Advanced PDF processing will be limited.")
    PDFPLUMBER_AVAILABLE = False

# Word processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available. Word document processing will be limited.")
    DOCX_AVAILABLE = False

# Image processing
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    logger.warning("PIL/pytesseract not available. OCR functionality will be limited.")
    OCR_AVAILABLE = False

# Table extraction
try:
    import warnings
    # Suppress jpype warnings since subprocess fallback works fine
    with warnings.catch_warnings():
        warnings.filterwarnings("ignore", message="Failed to import jpype dependencies")
        import tabula
    TABULA_AVAILABLE = True
    logger.info("Tabula-py loaded successfully")
except ImportError as e:
    logger.warning(f"Tabula not available: {e}. PDF table extraction will use pdfplumber only.")
    tabula = None
    TABULA_AVAILABLE = False

# Suppress jpype warnings if using subprocess fallback
import warnings
warnings.filterwarnings("ignore", message="Failed to import jpype dependencies.*")

class DocumentProcessor:
    """Processor for extracting financial data from various document formats."""
    
    def __init__(self, config: Any):
        """Initialize document processor."""
        self.config = config
        self.supported_formats = ['.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg']
    
    def process_file(self, uploaded_file) -> Dict[str, Any]:
        """Process an uploaded file and extract financial data."""
        try:
            # Get file extension
            file_extension = Path(uploaded_file.name).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Read file content
            file_content = uploaded_file.read()
            
            # Check if file is empty
            if not file_content or len(file_content) == 0:
                raise ValueError(f"File {uploaded_file.name} appears to be empty")
            
            logger.info(f"Processing file: {uploaded_file.name} ({len(file_content)} bytes)")
            
            # Process based on file type
            if file_extension == '.pdf':
                result = self._process_pdf(file_content, uploaded_file.name)
            elif file_extension in ['.docx', '.doc']:
                result = self._process_word(file_content, uploaded_file.name)
            elif file_extension in ['.png', '.jpg', '.jpeg']:
                result = self._process_image(file_content, uploaded_file.name)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Validate that we extracted some content
            self._validate_extracted_content(result)
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {uploaded_file.name}: {e}")
            raise
    
    def _process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF files and extract financial data."""
        result = {
            'filename': filename,
            'type': 'pdf',
            'text': '',
            'tables': [],
            'financial_data': {}
        }
        
        if not PYPDF2_AVAILABLE and not PDFPLUMBER_AVAILABLE:
            logger.error("No PDF processing libraries available. Please install PyPDF2 or pdfplumber.")
            result['text'] = "PDF processing unavailable - missing dependencies"
            return result
        
        try:
            text_content = []
            
            # Try PyPDF2 first if available
            if PYPDF2_AVAILABLE:
                try:
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    logger.info(f"PDF has {len(pdf_reader.pages)} pages")
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_content.append(page_text)
                                logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                            else:
                                logger.warning(f"Page {page_num + 1} appears to contain no extractable text (might be image-based)")
                        except Exception as e:
                            logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                            continue
                except Exception as e:
                    logger.error(f"PyPDF2 processing failed: {str(e)}")
            
            # Fallback to basic text if no content extracted
            if not text_content:
                text_content = ["PDF content could not be extracted. Please ensure the PDF contains text-based content."]
            
            result['text'] = '\n'.join(text_content)
            
            # If no text was extracted, this might be a scanned PDF
            if not result['text'].strip():
                logger.warning(f"No text extracted from PDF {filename}. This might be a scanned document.")
                
            # Extract tables using pdfplumber if available
            if PDFPLUMBER_AVAILABLE:
                try:
                    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                        for page_num, page in enumerate(pdf.pages):
                            try:
                                tables = page.extract_tables()
                                for table_num, table in enumerate(tables):
                                    if table and len(table) > 1:  # Skip empty tables
                                        # Clean table data
                                        cleaned_table = []
                                        for row in table:
                                            if row and any(cell for cell in row if cell is not None and str(cell).strip()):
                                                cleaned_table.append([str(cell).strip() if cell else '' for cell in row])
                                        
                                        if cleaned_table:
                                            df = pd.DataFrame(cleaned_table[1:], columns=cleaned_table[0])
                                            result['tables'].append(df)
                                            logger.debug(f"Extracted table {table_num + 1} from page {page_num + 1}")
                            except Exception as e:
                                logger.warning(f"Failed to extract tables from page {page_num + 1}: {e}")
                except Exception as e:
                    logger.warning(f"pdfplumber table extraction failed: {str(e)}")
            
            # Try tabula for better table extraction (if available)
            if TABULA_AVAILABLE:
                try:
                    # Suppress tabula/jpype warnings
                    import warnings
                    with warnings.catch_warnings():
                        warnings.simplefilter("ignore")
                        
                        # Save to temporary file for tabula
                        temp_file = f"temp_{filename}"
                        with open(temp_file, 'wb') as f:
                            f.write(file_content)
                        
                        # Use subprocess mode to avoid jpype issues
                        import os
                        old_java_options = os.environ.get('JAVA_TOOL_OPTIONS', '')
                        os.environ['JAVA_TOOL_OPTIONS'] = old_java_options + ' -Djava.awt.headless=true'
                        
                        tabula_tables = tabula.read_pdf(
                            temp_file, 
                            pages='all', 
                            multiple_tables=True,
                            silent=True,  # Suppress tabula output
                            java_options='-Djava.awt.headless=true'
                        )
                        
                        if tabula_tables:
                            result['tables'].extend(tabula_tables)
                            logger.info(f"Tabula extracted {len(tabula_tables)} additional tables")
                        
                        # Clean up temp file
                        Path(temp_file).unlink()
                        
                        # Restore java options
                        if old_java_options:
                            os.environ['JAVA_TOOL_OPTIONS'] = old_java_options
                        else:
                            os.environ.pop('JAVA_TOOL_OPTIONS', None)
                    
                except Exception as e:
                    logger.warning(f"Tabula extraction failed: {e}")
            else:
                logger.info("Tabula not available, using pdfplumber for table extraction only")
            
            # Extract financial data
            result['financial_data'] = self._extract_financial_data(result['text'], result['tables'])
            
        except Exception as e:
            logger.error(f"Error processing PDF {filename}: {e}")
            raise
        
        return result
    
    def _process_word(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process Word documents and extract financial data."""
        result = {
            'filename': filename,
            'type': 'word',
            'text': '',
            'tables': [],
            'financial_data': {}
        }
        
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available. Cannot process Word documents.")
            result['text'] = "Word document processing unavailable - missing python-docx dependency"
            return result
        
        try:
            # Extract text and tables from Word document
            doc = Document(io.BytesIO(file_content))
            
            # Extract text
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            result['text'] = '\\n'.join(text_content)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                    result['tables'].append(df)
            
            # Extract financial data
            result['financial_data'] = self._extract_financial_data(result['text'], result['tables'])
            
        except Exception as e:
            logger.error(f"Error processing Word document {filename}: {e}")
            raise
        
        return result
    
    def _process_image(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process image files using OCR and extract financial data."""
        result = {
            'filename': filename,
            'type': 'image',
            'text': '',
            'tables': [],
            'financial_data': {}
        }
        
        try:
            # Open image
            image = Image.open(io.BytesIO(file_content))
            
            # Perform OCR
            ocr_text = pytesseract.image_to_string(image)
            result['text'] = ocr_text
            
            # Try to extract table structure from OCR text
            table_data = self._parse_ocr_tables(ocr_text)
            if table_data:
                result['tables'] = table_data
            
            # Extract financial data
            result['financial_data'] = self._extract_financial_data(result['text'], result['tables'])
            
        except Exception as e:
            logger.error(f"Error processing image {filename}: {e}")
            raise
        
        return result
    
    def _parse_ocr_tables(self, ocr_text: str) -> List[pd.DataFrame]:
        """Parse table structure from OCR text."""
        tables = []
        
        try:
            # Split text into lines
            lines = ocr_text.split('\\n')
            
            # Look for table-like patterns
            table_lines = []
            for line in lines:
                # Check if line contains multiple numbers or currency values
                if self._is_table_line(line):
                    table_lines.append(line)
            
            if table_lines:
                # Convert to DataFrame
                table_data = []
                for line in table_lines:
                    # Split by whitespace or tabs
                    row = re.split(r'\s{2,}|\t', line.strip())
                    if len(row) > 1:
                        table_data.append(row)
                
                if table_data:
                    df = pd.DataFrame(table_data)
                    tables.append(df)
        
        except Exception as e:
            logger.warning(f"Error parsing OCR tables: {e}")
        
        return tables
    
    def _is_table_line(self, line: str) -> bool:
        """Check if a line appears to be part of a financial table."""
        # Look for patterns like currency values, percentages, numbers
        currency_pattern = r'\$[\d,]+(?:\.\d{2})?'
        percentage_pattern = r'\d+(?:\.\d+)?%'
        number_pattern = r'\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\b'
        
        currency_count = len(re.findall(currency_pattern, line))
        percentage_count = len(re.findall(percentage_pattern, line))
        number_count = len(re.findall(number_pattern, line))
        
        # Consider it a table line if it has multiple financial values
        return (currency_count + percentage_count + number_count) >= 2
    
    def _extract_financial_data(self, text: str, tables: List[pd.DataFrame]) -> Dict[str, Any]:
        """Extract financial metrics and data from text and tables."""
        financial_data = {}
        
        try:
            # Extract common financial metrics from text
            financial_data.update(self._extract_metrics_from_text(text))
            
            # Extract data from tables
            financial_data.update(self._extract_metrics_from_tables(tables))
            
        except Exception as e:
            logger.error(f"Error extracting financial data: {e}")
        
        return financial_data
    
    def _extract_metrics_from_text(self, text: str) -> Dict[str, Any]:
        """Extract financial metrics from text using regex patterns."""
        metrics = {}
        
        # Common financial metric patterns
        patterns = {
            'revenue': r'(?:revenue|sales|turnover)[\s:]*\$?([\d,]+(?:\.\d+)?)',
            'net_income': r'(?:net income|net profit|net earnings)[\s:]*\$?([\d,]+(?:\.\d+)?)',
            'total_assets': r'(?:total assets)[\s:]*\$?([\d,]+(?:\.\d+)?)',
            'total_equity': r'(?:total equity|shareholders\'? equity)[\s:]*\$?([\d,]+(?:\.\d+)?)',
            'gross_profit': r'(?:gross profit)[\s:]*\$?([\d,]+(?:\.\d+)?)',
            'operating_income': r'(?:operating income|operating profit)[\s:]*\$?([\d,]+(?:\.\d+)?)',
        }
        
        for metric_name, pattern in patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                try:
                    # Take the first match and convert to float
                    value = float(matches[0].replace(',', ''))
                    metrics[metric_name] = value
                except ValueError:
                    pass
        
        return metrics
    
    def _extract_metrics_from_tables(self, tables: List[pd.DataFrame]) -> Dict[str, Any]:
        """Extract financial metrics from structured tables."""
        metrics = {}
        
        for table in tables:
            try:
                # Look for financial statement patterns
                if self._is_income_statement(table):
                    metrics.update(self._parse_income_statement(table))
                elif self._is_balance_sheet(table):
                    metrics.update(self._parse_balance_sheet(table))
                elif self._is_cash_flow_statement(table):
                    metrics.update(self._parse_cash_flow_statement(table))
            
            except Exception as e:
                logger.warning(f"Error parsing table: {e}")
                continue
        
        return metrics
    
    def _is_income_statement(self, table: pd.DataFrame) -> bool:
        """Check if table appears to be an income statement."""
        if table.empty:
            return False
        
        text_content = ' '.join(table.astype(str).values.flatten()).lower()
        income_keywords = ['revenue', 'sales', 'net income', 'gross profit', 'operating income']
        
        return sum(keyword in text_content for keyword in income_keywords) >= 2
    
    def _is_balance_sheet(self, table: pd.DataFrame) -> bool:
        """Check if table appears to be a balance sheet."""
        if table.empty:
            return False
        
        text_content = ' '.join(table.astype(str).values.flatten()).lower()
        balance_keywords = ['assets', 'liabilities', 'equity', 'current assets', 'total assets']
        
        return sum(keyword in text_content for keyword in balance_keywords) >= 2
    
    def _is_cash_flow_statement(self, table: pd.DataFrame) -> bool:
        """Check if table appears to be a cash flow statement."""
        if table.empty:
            return False
        
        text_content = ' '.join(table.astype(str).values.flatten()).lower()
        cash_keywords = ['cash flow', 'operating activities', 'investing activities', 'financing activities']
        
        return sum(keyword in text_content for keyword in cash_keywords) >= 2
    
    def _parse_income_statement(self, table: pd.DataFrame) -> Dict[str, Any]:
        """Parse income statement data from table."""
        metrics = {}
        
        # Implementation for parsing income statement
        # This would involve looking for specific line items and extracting values
        
        return metrics
    
    def _parse_balance_sheet(self, table: pd.DataFrame) -> Dict[str, Any]:
        """Parse balance sheet data from table."""
        metrics = {}
        
        # Implementation for parsing balance sheet
        # This would involve looking for specific line items and extracting values
        
        return metrics
    
    def _parse_cash_flow_statement(self, table: pd.DataFrame) -> Dict[str, Any]:
        """Parse cash flow statement data from table."""
        metrics = {}
        
        # Implementation for parsing cash flow statement
        # This would involve looking for specific line items and extracting values
        
        return metrics
    
    def _validate_extracted_content(self, result: Dict[str, Any]) -> None:
        """Validate that we extracted meaningful content from the document."""
        filename = result.get('filename', 'Unknown')
        text = result.get('text', '')
        tables = result.get('tables', [])
        
        # Check if we have any text content
        if not text or len(text.strip()) < 10:
            if not tables or len(tables) == 0:
                logger.warning(f"No meaningful content extracted from {filename}")
                raise ValueError(f"Could not extract any readable content from {filename}. "
                               f"The file might be encrypted, image-based, or corrupted.")
        
        logger.info(f"Successfully extracted {len(text)} characters and {len(tables)} tables from {filename}")
